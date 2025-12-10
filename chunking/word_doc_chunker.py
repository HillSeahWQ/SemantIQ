"""
Word Document Chunker with intelligent content analysis and processing.

This chunker:
- Respects document structure (sections, headings, paragraphs)
- Extracts and processes tables, images, and diagrams
- Uses semantic chunking based on document structure
- Provides rich, unified metadata compatible with other document types

Key Features:

1. Structure-Aware Chunking
   - Chunks by sections (defined by heading styles)
   - Keeps related content together (headings + their content)
   - Handles nested heading levels intelligently

2. Multimodal Content Extraction
   - Tables: Extracted as markdown format
   - Images: Processed with vision model for descriptions
   - Diagrams: Handled as images with contextual analysis

3. Smart Chunking Strategy
   - Sections with text: Chunked together with their heading
   - Tables: Separate chunks with surrounding context
   - Images: Separate chunks with captions and descriptions
   - Size management: Large sections split at paragraph boundaries

4. Unified Metadata Format
   - Universal fields: source_file, chunk_id, page_number, chunk_type
   - Document-specific fields stored in dedicated metadata dict
   - Compatible with existing ChunkMetadata structure

## Usage

```python
# Install dependencies
pip install python-docx langchain-openai pillow

# Set environment variables
export OPENAI_API_KEY="your-key"

# Use the chunker
chunker = WordDocumentChunker(
    max_chunk_size=1000,  # Maximum characters per chunk
    vision_model="gpt-4o",
    process_images=True
)

chunks = chunker.chunk("document.docx")

# Access chunks
for chunk in chunks:
    print(f"Type: {chunk.metadata.chunk_type.value}")
    print(f"Content: {chunk.content[:200]}...")
    print(f"Metadata: {chunk.metadata.to_dict()}")
```

## Chunking Strategy Details

1. **Section-based chunking** (primary strategy):
   - Each heading + its content forms a potential chunk
   - Respects document hierarchy (Heading 1, 2, 3, etc.)
   - Maintains semantic coherence

2. **Size-based splitting** (when needed):
   - Large sections split at paragraph boundaries
   - Preserves sentence integrity
   - Metadata tracks if chunk is part of larger section

3. **Special content handling**:
   - Tables: Always separate chunks (easier for table-specific retrieval)
   - Images: Separate chunks with vision processing
   - Lists: Kept with their parent section

## Metadata Schema

Universal fields (all document types):
- source_file: str
- chunk_id: int
- page_number: int (approximated for Word docs)
- chunk_type: ChunkType

Document-specific fields (in document_specific_metadata):
- heading_level: int
- heading_text: str
- section_number: str
- has_lists: bool
- paragraph_count: int
- style_name: str
- is_continuation: bool (if chunk was split from larger section)
- parent_section: str
"""

import base64
import io
import logging
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Optional
from collections import defaultdict

from docx import Document
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from PIL import Image
from langchain_openai import ChatOpenAI
from langchain_core.messages import HumanMessage

from .base import BaseChunker, Chunk, ChunkMetadata, ChunkType
from utils.logger import get_logger

logger = get_logger(__name__)


@dataclass
class WordChunkMetadata(ChunkMetadata):
    """Extended metadata for Word document chunks with unified format."""
    
    # Universal fields (compatible with all document types)
    page_number: int = 0  # Approximated based on content position
    chunk_type: ChunkType = ChunkType.TEXT
    
    # Document-specific metadata (Word-specific fields)
    document_specific_metadata: Dict[str, Any] = field(default_factory=dict)
    
    # Preview and content info (universal)
    text_length: int = 0
    preview: str = ""  # First 200 chars for quick reference


class WordDocumentChunker(BaseChunker):
    """
    Word document chunker with intelligent structure analysis.
    
    Features:
    - Structure-aware chunking (sections, headings, paragraphs)
    - Table and image extraction with vision processing
    - Unified metadata format compatible with other document types
    - Smart size management with semantic preservation
    """
    
    # Heading styles that define sections
    HEADING_STYLES = [
        'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 
        'Heading 5', 'Heading 6', 'Title', 'Subtitle'
    ]
    
    def __init__(
        self,
        max_chunk_size: int = 1000,
        vision_model: str = "gpt-4o",
        process_images: bool = True,
        log_level: str = "INFO"
    ):
        """
        Initialize the Word document chunker.
        
        Args:
            max_chunk_size: Maximum characters per chunk (soft limit)
            vision_model: Vision-capable LLM model name for image processing
            process_images: Whether to process images with vision model
            log_level: Logging level (DEBUG, INFO, WARNING, ERROR)
        """
        self.max_chunk_size = max_chunk_size
        self.vision_model = ChatOpenAI(model=vision_model) if process_images else None
        self.process_images = process_images
        logger.setLevel(getattr(logging, log_level.upper()))
        
    def chunk(self, file_path: str | Path) -> List[Chunk]:
        """Process Word document and generate chunks."""
        return self.chunk_document(str(file_path))
    
    def get_metadata_schema(self) -> Dict[str, type]:
        """Return metadata schema for automatic DB schema generation."""
        return {
            "source_file": str,
            "chunk_id": int,
            "page_number": int,
            "chunk_type": str,
            "text_length": int,
            "preview": str,
            "document_specific_metadata": str,  # Stored as JSON string
        }
    
    def chunk_document(self, doc_path: str) -> List[Chunk]:
        """
        Process Word document and generate chunks.
        
        Args:
            doc_path: Path to Word document file
            
        Returns:
            List of Chunk objects with content and metadata
        """
        start_time = datetime.now()
        logger.info(f"{'='*80}")
        logger.info(f"Starting Word document chunking: {doc_path}")
        logger.info(f"Max chunk size: {self.max_chunk_size:,} characters")
        logger.info(f"Image processing: {'Enabled' if self.process_images else 'Disabled'}")
        logger.info(f"{'='*80}")
        
        try:
            doc = Document(doc_path)
            logger.info(f"Successfully opened document")
            logger.info(f"Total paragraphs: {len(doc.paragraphs)}")
            logger.info(f"Total tables: {len(doc.tables)}")
        except Exception as e:
            logger.error(f"[ERROR] - Failed to open document: {e}")
            raise
        
        chunks = []
        stats = {
            "text_chunks": 0,
            "table_chunks": 0,
            "image_chunks": 0,
            "mixed_chunks": 0,
            "sections_processed": 0,
            "total_tables": 0,
            "total_images": 0
        }
        
        # Parse document structure
        sections = self._parse_document_structure(doc, doc_path)
        logger.info(f"Parsed document into {len(sections)} sections")
        
        # Process each section
        for section_idx, section in enumerate(sections):
            logger.info(f"")
            logger.info(f"{'─'*80}")
            logger.info(f"Processing section {section_idx + 1}/{len(sections)}: {section['heading'][:50]}...")
            
            try:
                section_chunks = self._process_section(doc_path, doc, section, section_idx)
                
                for chunk in section_chunks:
                    chunk.metadata.chunk_id = len(chunks)
                    chunks.append(chunk)
                    
                    # Update statistics
                    if chunk.metadata.chunk_type == ChunkType.TABLE:
                        stats["table_chunks"] += 1
                        stats["total_tables"] += 1
                    elif chunk.metadata.chunk_type == ChunkType.IMAGE_HEAVY_PAGE:
                        stats["image_chunks"] += 1
                        stats["total_images"] += 1
                    elif chunk.metadata.chunk_type == ChunkType.MIXED:
                        stats["mixed_chunks"] += 1
                    else:
                        stats["text_chunks"] += 1
                
                stats["sections_processed"] += 1
                logger.info(f"[SUCCESS] - Created {len(section_chunks)} chunk(s) from section")
                
            except Exception as e:
                logger.error(f"[ERROR] - Error processing section {section_idx + 1}: {e}")
                logger.exception("Full error details:")
                continue
        
        elapsed = (datetime.now() - start_time).total_seconds()
        
        logger.info(f"")
        logger.info(f"{'='*80}")
        logger.info(f"CHUNKING COMPLETE")
        logger.info(f"{'='*80}")
        logger.info(f"Total chunks created: {len(chunks)}")
        logger.info(f"Time elapsed: {elapsed:.2f}s")
        logger.info(f"")
        logger.info(f"Chunk Type Distribution:")
        logger.info(f"   • Text chunks: {stats['text_chunks']}")
        logger.info(f"   • Table chunks: {stats['table_chunks']}")
        logger.info(f"   • Image chunks: {stats['image_chunks']}")
        logger.info(f"   • Mixed chunks: {stats['mixed_chunks']}")
        logger.info(f"")
        logger.info(f"Content Statistics:")
        logger.info(f"   • Sections processed: {stats['sections_processed']}")
        logger.info(f"   • Total tables: {stats['total_tables']}")
        logger.info(f"   • Total images: {stats['total_images']}")
        logger.info(f"{'='*80}")
        
        return chunks
    
    def _parse_document_structure(self, doc: Document, doc_path: str) -> List[Dict[str, Any]]:
        """
        Parse document into sections based on heading structure.
        
        Returns list of section dictionaries with heading and content.
        """
        sections = []
        current_section = None
        
        # Get all elements (paragraphs and tables) in order
        elements = []
        for element in doc.element.body:
            if isinstance(element, CT_P):
                elements.append(('paragraph', Paragraph(element, doc)))
            elif isinstance(element, CT_Tbl):
                elements.append(('table', DocxTable(element, doc)))
        
        for elem_type, element in elements:
            if elem_type == 'paragraph':
                para = element
                
                # Check if this is a heading
                if self._is_heading(para):
                    # Save previous section if exists
                    if current_section is not None:
                        sections.append(current_section)
                    
                    # Start new section
                    current_section = {
                        'heading': para.text.strip() or "(Untitled Section)",
                        'heading_level': self._get_heading_level(para),
                        'style_name': para.style.name,
                        'content': [],
                        'tables': [],
                        'images': [],
                        'has_lists': False
                    }
                else:
                    # Add to current section
                    if current_section is None:
                        # Create initial section if document doesn't start with heading
                        current_section = {
                            'heading': "(Document Start)",
                            'heading_level': 0,
                            'style_name': "Normal",
                            'content': [],
                            'tables': [],
                            'images': [],
                            'has_lists': False
                        }
                    
                    if para.text.strip():
                        current_section['content'].append(para.text)
                        
                        # Check for lists
                        if para.style.name.startswith('List'):
                            current_section['has_lists'] = True
                    
                    # Check for inline images
                    if self._has_inline_images(para):
                        current_section['images'].append({
                            'type': 'inline',
                            'paragraph': para
                        })
            
            elif elem_type == 'table':
                if current_section is None:
                    current_section = {
                        'heading': "(Document Start)",
                        'heading_level': 0,
                        'style_name': "Normal",
                        'content': [],
                        'tables': [],
                        'images': [],
                        'has_lists': False
                    }
                
                current_section['tables'].append(element)
        
        # Add final section
        if current_section is not None:
            sections.append(current_section)
        
        # If no sections found, create one
        if not sections:
            sections.append({
                'heading': Path(doc_path).stem,
                'heading_level': 0,
                'style_name': "Normal",
                'content': [],
                'tables': [],
                'images': [],
                'has_lists': False
            })
        
        return sections
    
    def _is_heading(self, paragraph: Paragraph) -> bool:
        """Check if paragraph is a heading."""
        return paragraph.style.name in self.HEADING_STYLES
    
    def _get_heading_level(self, paragraph: Paragraph) -> int:
        """Get heading level from paragraph."""
        style_name = paragraph.style.name
        if 'Heading' in style_name:
            try:
                return int(style_name.split()[-1])
            except ValueError:
                return 0
        elif style_name == 'Title':
            return 1
        elif style_name == 'Subtitle':
            return 2
        return 0
    
    def _has_inline_images(self, paragraph: Paragraph) -> bool:
        """Check if paragraph contains inline images."""
        return len(paragraph._element.xpath('.//a:blip')) > 0
    
    def _process_section(
        self, 
        doc_path: str,
        doc: Document,
        section: Dict[str, Any], 
        section_idx: int
    ) -> List[Chunk]:
        """
        Process a section and generate chunks.
        
        Strategy:
        1. Create text chunk(s) from section content
        2. Create separate chunks for each table
        3. Create separate chunks for each image
        """
        chunks = []
        
        # Approximate page number based on section index
        approx_page = section_idx + 1
        
        # Process text content
        if section['content']:
            text_content = '\n\n'.join(section['content'])
            
            # If content is too large, split into smaller chunks
            if len(text_content) > self.max_chunk_size:
                text_chunks = self._split_large_content(
                    text_content, 
                    section['content'],
                    doc_path,
                    section,
                    approx_page
                )
                chunks.extend(text_chunks)
            else:
                # Create single chunk for section
                chunk = self._create_text_chunk(
                    doc_path,
                    section,
                    text_content,
                    approx_page,
                    is_continuation=False
                )
                chunks.append(chunk)
        
        # Process tables separately
        for table_idx, table in enumerate(section['tables']):
            table_chunk = self._create_table_chunk(
                doc_path,
                section,
                table,
                table_idx,
                approx_page
            )
            chunks.append(table_chunk)
        
        # Process images separately
        if self.process_images:
            for img_idx, img_info in enumerate(section['images']):
                img_chunk = self._create_image_chunk(
                    doc_path,
                    doc,
                    section,
                    img_info,
                    img_idx,
                    approx_page
                )
                if img_chunk:
                    chunks.append(img_chunk)
        
        return chunks
    
    def _split_large_content(
        self,
        text_content: str,
        paragraphs: List[str],
        doc_path: str,
        section: Dict[str, Any],
        approx_page: int
    ) -> List[Chunk]:
        """Split large content into smaller chunks at paragraph boundaries."""
        chunks = []
        current_chunk_text = []
        current_length = 0
        
        for para in paragraphs:
            para_length = len(para)
            
            if current_length + para_length > self.max_chunk_size and current_chunk_text:
                # Create chunk from accumulated paragraphs
                chunk_text = '\n\n'.join(current_chunk_text)
                chunk = self._create_text_chunk(
                    doc_path,
                    section,
                    chunk_text,
                    approx_page,
                    is_continuation=len(chunks) > 0
                )
                chunks.append(chunk)
                
                # Start new chunk
                current_chunk_text = [para]
                current_length = para_length
            else:
                current_chunk_text.append(para)
                current_length += para_length
        
        # Add remaining content
        if current_chunk_text:
            chunk_text = '\n\n'.join(current_chunk_text)
            chunk = self._create_text_chunk(
                doc_path,
                section,
                chunk_text,
                approx_page,
                is_continuation=len(chunks) > 0
            )
            chunks.append(chunk)
        
        return chunks
    
    def _create_text_chunk(
        self,
        doc_path: str,
        section: Dict[str, Any],
        content: str,
        approx_page: int,
        is_continuation: bool
    ) -> Chunk:
        """Create a text chunk with unified metadata."""
        
        # Add heading to content if not a continuation
        if not is_continuation and section['heading'] != "(Document Start)":
            full_content = f"# {section['heading']}\n\n{content}"
        else:
            full_content = content
        
        # Document-specific metadata
        doc_specific = {
            'heading_level': section['heading_level'],
            'heading_text': section['heading'],
            'style_name': section['style_name'],
            'has_lists': section['has_lists'],
            'paragraph_count': len(content.split('\n\n')),
            'is_continuation': is_continuation,
            'parent_section': section['heading'] if is_continuation else None
        }
        
        metadata = WordChunkMetadata(
            source_file=doc_path,
            page_number=approx_page,
            chunk_type=ChunkType.TEXT,
            text_length=len(full_content),
            preview=full_content[:200],
            document_specific_metadata=doc_specific
        )
        
        return Chunk(content=full_content, metadata=metadata)
    
    def _create_table_chunk(
        self,
        doc_path: str,
        section: Dict[str, Any],
        table: DocxTable,
        table_idx: int,
        approx_page: int
    ) -> Chunk:
        """Create a chunk for a table."""
        
        # Convert table to markdown
        table_content = self._table_to_markdown(table)
        
        # Add context
        full_content = f"# {section['heading']}\n\n## Table {table_idx + 1}\n\n{table_content}"
        
        # Document-specific metadata
        doc_specific = {
            'heading_level': section['heading_level'],
            'heading_text': section['heading'],
            'table_index': table_idx,
            'table_rows': len(table.rows),
            'table_cols': len(table.columns),
            'parent_section': section['heading']
        }
        
        metadata = WordChunkMetadata(
            source_file=doc_path,
            page_number=approx_page,
            chunk_type=ChunkType.TABLE,
            text_length=len(full_content),
            preview=full_content[:200],
            document_specific_metadata=doc_specific
        )
        
        return Chunk(content=full_content, metadata=metadata)
    
    def _table_to_markdown(self, table: DocxTable) -> str:
        """Convert Word table to markdown format."""
        lines = []
        
        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
            
            # Add header separator after first row
            if row_idx == 0:
                lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
        
        return "\n".join(lines)
    
    def _create_image_chunk(
        self,
        doc_path: str,
        doc: Document,
        section: Dict[str, Any],
        img_info: Dict[str, Any],
        img_idx: int,
        approx_page: int
    ) -> Optional[Chunk]:
        """Create a chunk for an image with vision processing."""
        
        try:
            # Get image description using vision model
            paragraph = img_info['paragraph']
            description = self._process_image_with_vision(paragraph, doc, section['heading'])
            
            # Create content with context
            full_content = f"# {section['heading']}\n\n## Image {img_idx + 1}\n\n{description}"
            
            # Document-specific metadata
            doc_specific = {
                'heading_level': section['heading_level'],
                'heading_text': section['heading'],
                'image_index': img_idx,
                'image_type': img_info['type'],
                'parent_section': section['heading'],
                'is_vision_processed': True
            }
            
            metadata = WordChunkMetadata(
                source_file=doc_path,
                page_number=approx_page,
                chunk_type=ChunkType.IMAGE_HEAVY_PAGE,
                text_length=len(full_content),
                preview=full_content[:200],
                document_specific_metadata=doc_specific
            )
            
            return Chunk(content=full_content, metadata=metadata)
            
        except Exception as e:
            logger.warning(f"Failed to process image: {e}")
            return None
    
    def _extract_image_from_paragraph(self, paragraph: Paragraph, doc: Document) -> Optional[bytes]:
        """Extract image bytes from a paragraph."""
        try:
            # Find all blips (images) in the paragraph
            inline_shapes = paragraph._element.xpath('.//a:blip')
            if not inline_shapes:
                return None
            
            # Get first image's relationship ID
            blip = inline_shapes[0]
            rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            
            if not rId:
                return None
            
            # Get the image from document relationships
            image_part = doc.part.related_parts[rId]
            image_bytes = image_part.blob
            
            return image_bytes
            
        except Exception as e:
            logger.warning(f"Failed to extract image bytes: {e}")
            return None
    
    def _image_to_base64(self, image_bytes: bytes) -> str:
        """Convert image bytes to base64-encoded string."""
        try:
            # Open image with PIL to ensure it's valid and convert to PNG
            img = Image.open(io.BytesIO(image_bytes))
            
            # Convert to RGB if necessary (handles CMYK, grayscale, etc.)
            if img.mode not in ('RGB', 'RGBA'):
                img = img.convert('RGB')
            
            # Save as PNG to buffer
            buffer = io.BytesIO()
            img.save(buffer, format="PNG")
            
            # Encode to base64
            return base64.b64encode(buffer.getvalue()).decode("utf-8")
            
        except Exception as e:
            logger.error(f"Failed to convert image to base64: {e}")
            raise
    
    def _process_image_with_vision(self, paragraph: Paragraph, doc: Document, context: str) -> str:
        """Process an image with vision model following PDF chunker approach."""
        
        if not self.vision_model:
            return "[Image: Vision processing disabled]"
        
        logger.info(f"Vision processing initiated for image in section: {context}")
        logger.debug(f"    Extracting image from document...")
        
        try:
            # Extract image bytes
            image_bytes = self._extract_image_from_paragraph(paragraph, doc)
            if not image_bytes:
                return "[Image: Could not extract image data]"
            
            logger.debug(f"    Converting image to base64...")
            base64_image = self._image_to_base64(image_bytes)
            
            logger.debug(f"    Sending to vision model ({self.vision_model.model_name})...")
            
            # Use the same prompt structure as PDF chunker
            prompt = f"""Analyze this image from a document section titled "{context}" and provide a complete description.

Your output should include:

1. ALL TEXT CONTENT: Transcribe every piece of text visible in the image, maintaining the reading order and structure.
2. VISUAL ELEMENTS: Describe all charts, diagrams, photos, or visual content in detail.
3. LAYOUT AND STRUCTURE: Note how the content is organized within the image.
4. TABLES OR STRUCTURED DATA (if present): Describe the table structure and key data points.
5. CONTEXT: How this image relates to the document section "{context}".

OUTPUT FORMAT: Provide a flowing, readable description that captures all information in the image."""

            message = HumanMessage(
                content=[
                    {"type": "text", "text": prompt},
                    {
                        "type": "image_url",
                        "image_url": {"url": f"data:image/png;base64,{base64_image}"},
                    },
                ],
            )
            
            response = self.vision_model.invoke([message])
            logger.info(f"Vision processing complete ({len(response.content)} chars generated)")
            return response.content
            
        except Exception as e:
            logger.error(f"[ERROR] - Vision processing failed: {e}")
            return f"[Image in section: {context} - Processing failed: {str(e)}]"